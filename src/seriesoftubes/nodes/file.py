"""File ingestion and output node executor implementation"""

import csv
import glob
import io
import json
import random
from pathlib import Path
from typing import Any

import yaml
from pydantic import ValidationError

from seriesoftubes.file_security import (
    FileAccessMode,
    FileSecurityConfig,
    SecureFilePath,
)
from seriesoftubes.models import FileNodeConfig, Node
from seriesoftubes.nodes.base import NodeContext, NodeExecutor, NodeResult
from seriesoftubes.schemas import FileNodeInput, FileNodeOutput
from seriesoftubes.storage import StorageError, get_storage_backend
from seriesoftubes.template_engine import TemplateSecurityLevel, render_template

# Optional imports for document processing
try:
    import PyPDF2

    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl

    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    from bs4 import BeautifulSoup

    HAS_HTML = True
except ImportError:
    HAS_HTML = False


class FileNodeExecutor(NodeExecutor):
    """Executor for file ingestion and output nodes"""

    input_schema_class = FileNodeInput
    output_schema_class = FileNodeOutput
    
    def __init__(self, file_security_config: FileSecurityConfig | None = None):
        """Initialize with optional file security configuration.
        
        Args:
            file_security_config: Security configuration for file access.
                                If None, file security is disabled (backward compatibility)
        """
        super().__init__()
        self.file_security = SecureFilePath(file_security_config) if file_security_config else None

    async def execute(self, node: Node, context: NodeContext) -> NodeResult:
        """Execute a file node"""
        if not isinstance(node.config, FileNodeConfig):
            return NodeResult(
                output=None,
                success=False,
                error=f"Invalid config type for file node: {type(node.config)}",
            )

        config = node.config

        try:
            # Prepare context for template rendering
            context_data = self.prepare_context_data(node, context)

            if config.mode == "write":
                return await self._execute_write_mode(node, config, context_data)
            else:
                return await self._execute_read_mode(node, config, context_data)

        except Exception as e:
            return NodeResult(
                output=None,
                success=False,
                error=str(e),
            )

    async def _execute_write_mode(
        self, node: Node, config: FileNodeConfig, context_data: dict[str, Any]
    ) -> NodeResult:
        """Execute file node in write mode - save data to storage"""
        try:
            # Get data to write from node inputs or context mapping
            # First check if there's a specific "data" field in inputs
            inputs = context_data.get("inputs", {})
            if isinstance(inputs, dict) and "data" in inputs:
                data = inputs["data"]
            else:
                # Otherwise use the entire inputs
                data = inputs
            
            # Render the write key
            write_key = self._render_template(config.write_key, context_data)
            
            # Add storage prefix if specified
            if config.storage_prefix:
                write_key = f"{config.storage_prefix.rstrip('/')}/{write_key}"
            
            # Determine content type based on format
            content_type = self._get_content_type(config.format_type)
            
            # Serialize data based on format
            content = self._serialize_data(data, config.format_type)
            
            if config.storage_type == "object":
                # Write to object storage
                storage = get_storage_backend()
                await storage.initialize()
                
                # Upload to storage
                stored_file = await storage.upload(
                    key=write_key,
                    content=content.encode(config.encoding) if isinstance(content, str) else content,
                    content_type=content_type,
                    metadata={
                        "node_name": node.name,
                        "workflow_execution_id": context_data.get("execution_id", "unknown"),
                    }
                )
                
                # Get a pre-signed URL for the file
                file_url = await storage.get_url(write_key, expires_in=86400)  # 24 hours
                
                output = {
                    "storage_type": "object",
                    "key": write_key,
                    "size": stored_file.size,
                    "url": file_url,
                    "content_type": content_type,
                }
            else:
                # Write to local filesystem
                path = Path(write_key)
                
                # Validate path with security if enabled
                if self.file_security:
                    path = self.file_security.validate_path(
                        path,
                        mode=FileAccessMode.WRITE,
                        must_exist=False
                    )
                
                # Ensure parent directory exists
                path.parent.mkdir(parents=True, exist_ok=True)
                
                # Write file
                if isinstance(content, bytes):
                    path.write_bytes(content)
                else:
                    path.write_text(content, encoding=config.encoding)
                
                output = {
                    "storage_type": "local",
                    "path": str(path),
                    "size": path.stat().st_size,
                }
            
            # Log file write for security auditing
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"File write: Wrote to '{write_key}' (storage: {config.storage_type})")
            
            return NodeResult(
                output=output,
                success=True,
                metadata={
                    "mode": "write",
                    "storage_type": config.storage_type,
                    "key": write_key,
                },
            )
            
        except Exception as e:
            return NodeResult(
                output=None,
                success=False,
                error=f"Write failed: {str(e)}",
            )

    async def _execute_read_mode(
        self, node: Node, config: FileNodeConfig, context_data: dict[str, Any]
    ) -> NodeResult:
        """Execute file node in read mode - read data from storage"""
        try:
            # Always validate input when schema is defined
            # Extract the rendered path/pattern for validation
            rendered_path = None
            rendered_pattern = None

            if config.path:
                rendered_path = self._render_template(config.path, context_data)
            elif config.pattern:
                rendered_pattern = self._render_template(config.pattern, context_data)

            input_data = {
                "path": rendered_path,
                "pattern": rendered_pattern,
            }

            try:
                validated_input = self.validate_input(input_data)
                # Update config with validated values
                if validated_input.get("path"):
                    config.path = validated_input["path"]
                if validated_input.get("pattern"):
                    config.pattern = validated_input["pattern"]
            except ValidationError as e:
                # Format validation errors for clarity
                error_details = []
                for error in e.errors():
                    field = ".".join(str(x) for x in error["loc"])
                    error_details.append(f"  - {field}: {error['msg']}")

                return NodeResult(
                    output=None,
                    success=False,
                    error=f"Input validation failed for node '{node.name}':\n"
                    + "\n".join(error_details),
                )

            # Get file paths based on storage type
            if config.storage_type == "object":
                paths = await self._get_object_storage_paths(config, context_data)
            else:
                paths = self._get_file_paths(config, context_data)

            if not paths:
                return NodeResult(
                    output=None,
                    success=False,
                    error="No files found matching criteria",
                )

            # Process based on number of files and output mode
            if len(paths) == 1:
                # Single file
                data = await self._process_single_file(paths[0], config)
            else:
                # Multiple files
                data = await self._process_multiple_files(paths, config)

            # Structure the output
            output = {
                "data": data,
                "metadata": {
                    "files_read": len(paths),
                    "output_mode": config.output_mode,
                    "format": config.format_type,
                    "storage_type": config.storage_type,
                },
            }

            # Always validate output when schema is defined
            try:
                output = self.validate_output(output)
            except ValidationError as e:
                # Format validation errors for clarity
                error_details = []
                for error in e.errors():
                    field = ".".join(str(x) for x in error["loc"])
                    error_details.append(f"  - {field}: {error['msg']}")

                return NodeResult(
                    output=None,
                    success=False,
                    error=f"Output validation failed for node '{node.name}':\n"
                    + "\n".join(error_details),
                )

            return NodeResult(
                output=output,
                success=True,
                metadata={
                    "files_read": len(paths),
                    "output_mode": config.output_mode,
                    "storage_type": config.storage_type,
                },
            )

        except Exception as e:
            return NodeResult(
                output=None,
                success=False,
                error=str(e),
            )

    async def _get_object_storage_paths(
        self, config: FileNodeConfig, context: dict[str, Any]
    ) -> list[dict[str, str]]:
        """Get list of files from object storage"""
        paths = []
        
        try:
            storage = get_storage_backend()
            await storage.initialize()
            
            if config.path:
                # Single file path
                rendered_path = self._render_template(config.path, context)
                
                # Add storage prefix if specified
                if config.storage_prefix:
                    rendered_path = f"{config.storage_prefix.rstrip('/')}/{rendered_path}"
                
                # Check if file exists
                files = await storage.list(prefix=rendered_path, max_keys=1)
                if files and files[0].key == rendered_path:
                    paths.append({
                        "key": rendered_path,
                        "storage_type": "object"
                    })
                elif not config.skip_errors:
                    raise FileNotFoundError(f"File not found in storage: {rendered_path}")
                    
            elif config.pattern:
                # Pattern matching in object storage
                rendered_pattern = self._render_template(config.pattern, context)
                
                # Add storage prefix if specified
                if config.storage_prefix:
                    rendered_pattern = f"{config.storage_prefix.rstrip('/')}/{rendered_pattern}"
                
                # Extract prefix from pattern for efficient listing
                prefix = rendered_pattern.split('*')[0].rstrip('/')
                
                # List files with prefix
                files = await storage.list(prefix=prefix, max_keys=1000)
                
                # Filter files matching the pattern
                import fnmatch
                for file_info in files:
                    if fnmatch.fnmatch(file_info.key, rendered_pattern):
                        paths.append({
                            "key": file_info.key,
                            "storage_type": "object"
                        })
                        
        except StorageError as e:
            if not config.skip_errors:
                raise ValueError(f"Storage error: {e}") from e
                
        return sorted(paths, key=lambda x: x["key"])

    def _get_file_paths(
        self, config: FileNodeConfig, context: dict[str, Any]
    ) -> list[dict[str, str]]:
        """Get list of file paths based on config"""
        paths = []

        if config.path:
            # Single file path
            rendered_path = self._render_template(config.path, context)
            try:
                if self.file_security:
                    # Validate path with security checks
                    path = self.file_security.validate_path(
                        rendered_path,
                        mode=FileAccessMode.READ,
                        must_exist=not config.skip_errors
                    )
                else:
                    # No security - just use the path
                    path = Path(rendered_path)
                    if not path.exists() and not config.skip_errors:
                        raise FileNotFoundError(f"File not found: {path}")
                
                if path.exists():
                    paths.append({
                        "path": str(path),
                        "storage_type": "local"
                    })
            except FileNotFoundError:
                if not config.skip_errors:
                    raise
            except Exception as e:
                if not config.skip_errors:
                    raise ValueError(f"Invalid file path: {e}") from e

        elif config.pattern:
            # Glob pattern
            rendered_pattern = self._render_template(config.pattern, context)
            
            # For glob patterns, we need to validate the base directory first
            # Extract the base directory from the pattern
            pattern_parts = rendered_pattern.split('/')
            base_parts = []
            for part in pattern_parts:
                if '*' in part or '?' in part or '[' in part:
                    break
                base_parts.append(part)
            
            # Handle absolute and relative paths
            if rendered_pattern.startswith('/'):
                base_dir = '/' + '/'.join(base_parts) if base_parts else '/'
            else:
                base_dir = '/'.join(base_parts) if base_parts else '.'
            
            try:
                if self.file_security:
                    # Validate base directory access
                    validated_base = self.file_security.validate_path(
                        base_dir,
                        mode=FileAccessMode.READ,
                        must_exist=True
                    )
                
                # Now perform glob
                found_paths = [Path(p) for p in glob.glob(rendered_pattern)]
                
                # Filter to only include files (not directories)
                found_paths = [p for p in found_paths if p.is_file()]
                
                # Validate each found path if security is enabled
                for p in found_paths:
                    try:
                        if self.file_security:
                            validated_path = self.file_security.validate_path(
                                p,
                                mode=FileAccessMode.READ,
                                must_exist=True
                            )
                            if validated_path.is_file():
                                paths.append({
                                    "path": str(validated_path),
                                    "storage_type": "local"
                                })
                        else:
                            # No security - just add the path
                            if p.is_file():
                                paths.append({
                                    "path": str(p),
                                    "storage_type": "local"
                                })
                    except Exception:
                        if not config.skip_errors:
                            raise
            except Exception as e:
                if not config.skip_errors:
                    raise ValueError(f"Invalid glob pattern or access denied: {e}") from e

        return sorted(paths, key=lambda x: x["path"])  # Consistent ordering

    async def _process_single_file(self, file_info: dict[str, str], config: FileNodeConfig) -> Any:
        """Process a single file"""
        # Log file access for security auditing
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"File access: Reading file '{file_info}' (format: {config.format_type})")
        
        content = await self._read_file(file_info, config)

        # Apply transformations
        if config.format_type in ["csv", "jsonl"] and isinstance(content, list):
            content = self._apply_filters(content, config)

        # Return based on output mode
        if config.output_mode == "list" and not isinstance(content, list):
            return [content]
        return content

    async def _process_multiple_files(
        self, file_infos: list[dict[str, str]], config: FileNodeConfig
    ) -> Any:
        """Process multiple files"""
        if config.merge:
            # Merge files into single output
            merged = []
            for file_info in file_infos:
                try:
                    content = await self._read_file(file_info, config)
                    if isinstance(content, list):
                        merged.extend(content)
                    else:
                        merged.append(content)
                except Exception:
                    if not config.skip_errors:
                        raise
                    # Skip this file

            # Apply filters to merged content
            if config.format_type in ["csv", "jsonl"]:
                merged = self._apply_filters(merged, config)

            return merged

        else:
            # Keep files separate
            output = {}
            for file_info in file_infos:
                try:
                    content = await self._read_file(file_info, config)
                    # Apply filters per file
                    if config.format_type in ["csv", "jsonl"] and isinstance(
                        content, list
                    ):
                        content = self._apply_filters(content, config)
                    
                    # Use path or key as the identifier
                    identifier = file_info.get("path") or file_info.get("key")
                    output[identifier] = content
                except Exception as e:
                    if not config.skip_errors:
                        raise
                    identifier = file_info.get("path") or file_info.get("key")
                    output[identifier] = {"error": str(e)}

            # Return based on output mode
            if config.output_mode == "list":
                return list(output.values())
            return output

    async def _read_file(self, file_info: dict[str, str], config: FileNodeConfig) -> Any:
        """Read a single file based on format"""
        if file_info["storage_type"] == "object":
            # Read from object storage
            storage = get_storage_backend()
            await storage.initialize()
            
            content_bytes = await storage.download(file_info["key"])
            
            # Create a file-like object for parsing
            if config.format_type in ["pdf"]:
                file_obj = io.BytesIO(content_bytes)
            else:
                # Decode text content
                content_str = content_bytes.decode(config.encoding)
                file_obj = io.StringIO(content_str)
            
            # Parse content based on format
            return self._parse_content(file_obj, file_info["key"], config)
        else:
            # Read from local filesystem
            path = Path(file_info["path"])
            return await self._read_local_file(path, config)

    async def _read_local_file(self, path: Path, config: FileNodeConfig) -> Any:
        """Read a single file from local filesystem"""
        # Auto-detect format from extension
        file_format = config.format_type
        if file_format == "auto":
            suffix = path.suffix.lower()
            if suffix in [".json"]:
                file_format = "json"
            elif suffix in [".jsonl", ".ndjson"]:
                file_format = "jsonl"
            elif suffix in [".yaml", ".yml"]:
                file_format = "yaml"
            elif suffix in [".csv"]:
                file_format = "csv"
            elif suffix in [".pdf"]:
                file_format = "pdf"
            elif suffix in [".docx"]:
                file_format = "docx"
            elif suffix in [".xlsx", ".xls"]:
                file_format = "xlsx"
            elif suffix in [".html", ".htm"]:
                file_format = "html"
            else:
                file_format = "txt"

        # Read based on format
        if file_format == "json":
            with open(path, encoding=config.encoding) as f:
                return json.load(f)
        elif file_format == "jsonl":
            return self._read_jsonl(path, config)
        elif file_format == "yaml":
            with open(path, encoding=config.encoding) as f:
                return yaml.safe_load(f)
        elif file_format == "csv":
            return self._read_csv(path, config)
        elif file_format == "pdf":
            return self._read_pdf(path, config)
        elif file_format == "docx":
            return self._read_docx(path, config)
        elif file_format == "xlsx":
            return self._read_xlsx(path, config)
        elif file_format == "html":
            return self._read_html(path, config)
        else:
            # Default to text
            with open(path, encoding=config.encoding) as f:
                return f.read()

    def _parse_content(self, file_obj: io.IOBase, filename: str, config: FileNodeConfig) -> Any:
        """Parse content from a file-like object"""
        # Auto-detect format from filename if needed
        file_format = config.format_type
        if file_format == "auto":
            suffix = Path(filename).suffix.lower()
            if suffix in [".json"]:
                file_format = "json"
            elif suffix in [".jsonl", ".ndjson"]:
                file_format = "jsonl"
            elif suffix in [".yaml", ".yml"]:
                file_format = "yaml"
            elif suffix in [".csv"]:
                file_format = "csv"
            elif suffix in [".pdf"]:
                file_format = "pdf"
            else:
                file_format = "txt"

        # Parse based on format
        if file_format == "json":
            return json.load(file_obj)
        elif file_format == "jsonl":
            return self._parse_jsonl(file_obj, config)
        elif file_format == "yaml":
            return yaml.safe_load(file_obj)
        elif file_format == "csv":
            return self._parse_csv(file_obj, config)
        elif file_format == "pdf":
            # PDF requires binary mode
            if not isinstance(file_obj, io.BytesIO):
                raise ValueError("PDF parsing requires binary content")
            return self._parse_pdf(file_obj, config)
        else:
            # Default to text
            return file_obj.read()

    def _parse_jsonl(self, file_obj: io.StringIO, config: FileNodeConfig) -> list[dict[str, Any]]:
        """Parse JSONL from file-like object"""
        rows = []
        for line_num, line in enumerate(file_obj):
            if line.strip():
                try:
                    rows.append(json.loads(line))
                except json.JSONDecodeError as e:
                    if not config.skip_errors:
                        msg = f"Invalid JSON on line {line_num + 1}: {e}"
                        raise ValueError(msg) from e
        return rows

    def _parse_csv(self, file_obj: io.StringIO, config: FileNodeConfig) -> list[dict[str, Any]]:
        """Parse CSV from file-like object"""
        rows = []
        if config.has_header:
            dict_reader = csv.DictReader(file_obj, delimiter=config.delimiter)
            for row in dict_reader:
                rows.append(dict(row))
        else:
            list_reader = csv.reader(file_obj, delimiter=config.delimiter)
            for row in list_reader:  # type: ignore[assignment]
                rows.append({f"col_{j}": val for j, val in enumerate(row)})
        return rows

    def _parse_pdf(self, file_obj: io.BytesIO, config: FileNodeConfig) -> Any:
        """Parse PDF from file-like object"""
        if not HAS_PDF:
            msg = "PDF support requires PyPDF2. Install with: pip install PyPDF2"
            raise ImportError(msg)

        if config.extract_text:
            # Extract text from PDF
            text = []
            pdf_reader = PyPDF2.PdfFileReader(file_obj)
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text.append(page.extractText())
            return "\n".join(text)
        else:
            # Return metadata only
            pdf_reader = PyPDF2.PdfFileReader(file_obj)
            return {
                "num_pages": pdf_reader.numPages,
                "metadata": pdf_reader.getDocumentInfo(),
            }

    def _read_jsonl(self, path: Path, config: FileNodeConfig) -> list[dict[str, Any]]:
        """Read JSONL file"""
        rows = []
        with open(path, encoding=config.encoding) as f:
            for line_num, line in enumerate(f):
                if line.strip():
                    try:
                        rows.append(json.loads(line))
                    except json.JSONDecodeError as e:
                        if not config.skip_errors:
                            msg = f"Invalid JSON on line {line_num + 1}: {e}"
                            raise ValueError(msg) from e
        return rows

    def _read_csv(self, path: Path, config: FileNodeConfig) -> list[dict[str, Any]]:
        """Read CSV file into list of dicts"""
        rows = []
        with open(path, encoding=config.encoding) as f:
            if config.has_header:
                dict_reader = csv.DictReader(f, delimiter=config.delimiter)
                for row in dict_reader:
                    rows.append(dict(row))
            else:
                list_reader = csv.reader(f, delimiter=config.delimiter)
                for row in list_reader:  # type: ignore[assignment]
                    rows.append({f"col_{j}": val for j, val in enumerate(row)})
        return rows

    def _apply_filters(self, records: list[Any], config: FileNodeConfig) -> list[Any]:
        """Apply sampling, limiting, etc. to records"""
        # Sample
        if config.sample is not None and 0 < config.sample < 1:
            num_samples = int(len(records) * config.sample)
            records = random.sample(records, num_samples)

        # Limit
        if config.limit is not None and config.limit > 0:
            records = records[: config.limit]

        return records

    def _read_pdf(self, path: Path, config: FileNodeConfig) -> Any:
        """Read PDF file"""
        if not HAS_PDF:
            msg = "PDF support requires PyPDF2. Install with: pip install PyPDF2"
            raise ImportError(msg)

        if config.extract_text:
            # Extract text from PDF
            text = []
            with open(path, "rb") as f:
                pdf_reader = PyPDF2.PdfFileReader(f)
                for page_num in range(pdf_reader.numPages):
                    page = pdf_reader.getPage(page_num)
                    text.append(page.extractText())
            return "\n".join(text)
        else:
            # Return metadata only
            with open(path, "rb") as f:
                pdf_reader = PyPDF2.PdfFileReader(f)
                return {
                    "num_pages": pdf_reader.numPages,
                    "metadata": pdf_reader.getDocumentInfo(),
                    "path": str(path),
                }

    def _read_docx(self, path: Path, config: FileNodeConfig) -> Any:
        """Read DOCX file"""
        if not HAS_DOCX:
            msg = (
                "DOCX support requires python-docx. Install with: "
                "pip install python-docx"
            )
            raise ImportError(msg)

        doc = Document(path)

        if config.extract_text:
            # Extract text from document
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        else:
            # Return structured content
            return {
                "paragraphs": [p.text for p in doc.paragraphs],
                "tables": [
                    [[cell.text for cell in row.cells] for row in table.rows]
                    for table in doc.tables
                ],
                "path": str(path),
            }

    def _read_xlsx(self, path: Path, config: FileNodeConfig) -> Any:
        """Read Excel file"""
        if not HAS_XLSX:
            msg = "Excel support requires openpyxl. Install with: pip install openpyxl"
            raise ImportError(msg)

        workbook = openpyxl.load_workbook(path, data_only=True)

        # For now, just read the first sheet as CSV-like data
        sheet = workbook.active
        if sheet is None:
            return []
        rows = []

        # Get all rows
        all_rows = list(sheet.iter_rows(values_only=True))

        if all_rows and config.has_header:
            # Use first row as headers
            headers = [str(h) if h else f"col_{i}" for i, h in enumerate(all_rows[0])]
            for row in all_rows[1:]:
                row_dict = {}
                for i, value in enumerate(row):
                    if i < len(headers):
                        row_dict[headers[i]] = value
                if any(row_dict.values()):  # Skip empty rows
                    rows.append(row_dict)
        else:
            # No headers
            for row in all_rows:
                row_dict = {f"col_{i}": value for i, value in enumerate(row)}
                if any(row_dict.values()):  # Skip empty rows
                    rows.append(row_dict)

        return rows

    def _read_html(self, path: Path, config: FileNodeConfig) -> Any:
        """Read HTML file"""
        if not HAS_HTML:
            msg = (
                "HTML support requires beautifulsoup4. Install with: "
                "pip install beautifulsoup4"
            )
            raise ImportError(msg)

        with open(path, encoding=config.encoding) as f:
            soup = BeautifulSoup(f, "html.parser")

        if config.extract_text:
            # Extract visible text
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text()

            # Break into lines and remove empty ones
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            return text
        else:
            # Return structured content
            return {
                "title": soup.title.string if soup.title else None,
                "text": soup.get_text(),
                "links": [
                    {"href": a.get("href"), "text": a.get_text()}
                    for a in soup.find_all("a")
                ],
                "path": str(path),
            }

    def _render_template(self, template_str: str, context: dict[str, Any]) -> str:
        """Render a template string securely"""
        # File paths should only use simple interpolation, no expressions
        return render_template(
            template_str, 
            context, 
            level=TemplateSecurityLevel.INTERPOLATION_ONLY,
            node_type="file"
        )

    def _serialize_data(self, data: Any, format_type: str) -> str | bytes:
        """Serialize data based on format type"""
        if format_type == "json":
            return json.dumps(data, indent=2)
        elif format_type == "jsonl":
            if isinstance(data, list):
                return "\n".join(json.dumps(item) for item in data)
            else:
                return json.dumps(data)
        elif format_type == "yaml":
            return yaml.dump(data, default_flow_style=False, sort_keys=False)
        elif format_type == "csv":
            if not isinstance(data, list) or not data:
                raise ValueError("CSV format requires a non-empty list of dictionaries")
            
            output = io.StringIO()
            writer = csv.DictWriter(output, fieldnames=data[0].keys())
            writer.writeheader()
            writer.writerows(data)
            return output.getvalue()
        else:
            # Default to text/string representation
            if isinstance(data, (dict, list)):
                return json.dumps(data, indent=2)
            else:
                return str(data)

    def _get_content_type(self, format_type: str) -> str:
        """Get content type for a format"""
        content_types = {
            "json": "application/json",
            "jsonl": "application/x-ndjson",
            "yaml": "text/yaml",
            "csv": "text/csv",
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "html": "text/html",
            "txt": "text/plain",
        }
        return content_types.get(format_type, "application/octet-stream")